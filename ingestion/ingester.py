from docx import Document

class IngestManager:
    def __init__(self, file_path: str):
        self.file_path = file_path


    def prepare_chunks_from_docx(self, chunk_size: int = 150, overlap: int = 50):
            """
            Вытаскивает текст из документа docx и разбивает его на чанки
            """
            text = self._extract_text_from_docx()
            chunks = self._chunk_text_with_context(text, chunk_size=chunk_size, overlap=overlap)
            return chunks


    def prepare_chunks_from_table(self):
        """
        Преобразует таблицу Word в список плоских чанков для ChromaDB.
        Каждый профиль создаёт отдельный документ с плоскими метаданными.
        На вход подается путь до файла.
        Предварительно надо отредактировать и положить информацию о форме обучения во второй столбец заголовка: очное, заочное, очно-заочное
        """
        doc = Document(self.file_path)
        
        chunks = []
        direction = None
        form = None

        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                
                # Пропускаем лишние строки
                if "всего" in values[1].lower():
                    continue
                
                # Определяем форму обучения
                if "очно-заочное" in values[1].lower():
                    form = "очно-заочная"
                    continue
                elif "заочное" in values[1].lower():
                    form = "заочная"
                    continue
                elif "очное" in values[1].lower():
                    form = "очная"
                    continue
                
                # Определяем major_id и secondary_id
                major_id, secondary_id = self._get_ids(values[0])

                # Определяем название направления
                direction, profile = self._get_direction_names(values[0], direction, values[1])

                # Сохраняем профиль и баллы
                if values[5]:
                    profile = profile
                    passing_budget = values[5]
                    passing_paid = values[9]

                    # Создаём отдельный документ для каждого профиля
                    chunk = {
                        "direction": direction,
                        "form": form,
                        "major_id": major_id,
                        "profile": profile,
                        "passing_budget": passing_budget,
                        "passing_paid": passing_paid,
                        "secondary_id": secondary_id
                    }
                    chunks.append(chunk)
        return chunks


    def _extract_text_from_docx(self) -> str:
        """
        Извлекает текст из docx файлов по абзацам
        """
        doc = Document(self.file_path)
        full_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # игнорируем пустые строки
                full_text.append(text)

        # Соединяем абзацы через двойной перенос для логики chunking
        return "\n\n".join(full_text)
    

    def _chunk_text_with_context(self, text: str, chunk_size: int = 150, overlap: int = 50) -> list[str]:
        """
        Функция разбиения текста на чанки с перекрытием, возвращает список чанков
        """
        paragraphs = text.split("\n\n")
        chunks = []

        for para in paragraphs:
            words = para.split()
            start = 0
            while start < len(words):
                end = start + chunk_size
                chunk = " ".join(words[start:end])
                chunks.append(chunk)
                start = max(end - overlap, end)  # перекрытие
        return chunks


    def _get_ids(self, current_id: str):
        """
        Определяет код направления подготовки
        """
        if len(current_id) > 8:
            major_id = current_id[:8]
            secondary_id = current_id
        else:
            major_id = current_id
            secondary_id = None
        return major_id, secondary_id


    def _get_direction_names(self, current_id: str, direction: str, name: str):
        """
        Определяет имена направления и профиля подготовки

        Пример:
        current_id = 09.03.01.01 ->
        major_name = Информатика и вычислительная техника
        profile_name = Веб-технологии
        """
        if len(current_id) <= 8:
            direction = name
            profile = ""
        else:
            profile = name
        return direction, profile